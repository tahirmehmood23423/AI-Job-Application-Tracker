"""
Text extraction from PDF and DOCX files.

Design notes:
- pdfplumber is the primary PDF extractor because it handles multi-column
  resumes better than pypdf. If it fails (rare but happens with malformed
  PDFs), we fall back to pypdf.
- For DOCX, python-docx is the standard. We extract paragraphs AND table
  cells, because many resumes use tables for layout (especially for
  side-by-side skills sections).
- Encrypted PDFs raise an error — we don't try to crack them.
- We strip the extracted text of obvious noise (repeated whitespace,
  page-number footers) before returning it, to save LLM tokens later.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

import pdfplumber
import pypdf
from docx import Document as DocxDocument

from app.core.exceptions import (
    EmptyFileError,
    TextExtractionError,
    UnsupportedFileTypeError,
)
from app.utils.logger import get_logger

logger = get_logger(__name__)


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

# A simple regex that matches lines like "Page 1 of 3" or just "1" centered
_PAGE_NUMBER_PATTERN = re.compile(r"^\s*(page\s+)?\d+\s*(of\s+\d+)?\s*$", re.IGNORECASE)


class TextExtractor:
    """Extracts plain text from a resume file (PDF or DOCX)."""

    def extract(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract text from the given file.

        Args:
            file_bytes: Raw bytes of the uploaded file.
            filename: Original filename, used only to detect the extension.

        Returns:
            Cleaned plain text from the resume.

        Raises:
            UnsupportedFileTypeError: Extension is not .pdf or .docx.
            EmptyFileError: The file produced no text.
            TextExtractionError: The file could not be parsed.
        """
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{ext}'. Allowed: {sorted(SUPPORTED_EXTENSIONS)}"
            )

        logger.info(
            "Extracting text from file",
            extra={"upload_filename": filename, "extension": ext, "size_bytes": len(file_bytes)},
        )
        if ext == ".pdf":
            text = self._extract_pdf(file_bytes)
        else:  # .docx
            text = self._extract_docx(file_bytes)

        cleaned = self._clean(text)

        if not cleaned or len(cleaned) < 50:
            # < 50 chars is almost certainly a failure (a scanned PDF with no
            # OCR, or an empty file). We don't try to OCR here — that's a
            # future enhancement.
            raise EmptyFileError(
                f"Extracted only {len(cleaned)} characters. "
                "The file may be empty, image-only (scanned), or corrupted."
            )

        logger.info("Text extracted", extra={"chars": len(cleaned)})
        return cleaned

    # ----- PDF -----

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Try pdfplumber first, fall back to pypdf."""
        try:
            return self._extract_pdf_with_pdfplumber(file_bytes)
        except Exception as e:
            logger.warning(
                "pdfplumber failed, falling back to pypdf",
                extra={"error": str(e)},
            )
            try:
                return self._extract_pdf_with_pypdf(file_bytes)
            except Exception as e2:
                raise TextExtractionError(
                    f"Could not extract text from PDF. "
                    f"pdfplumber error: {e}; pypdf error: {e2}"
                ) from e2

    def _extract_pdf_with_pdfplumber(self, file_bytes: bytes) -> str:
        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(
                    # Tweaked tolerances work better on resumes with tight line
                    # spacing.
                    x_tolerance=2,
                    y_tolerance=3,
                )
                if page_text:
                    parts.append(page_text)
        return "\n\n".join(parts)

    def _extract_pdf_with_pypdf(self, file_bytes: bytes) -> str:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        if reader.is_encrypted:
            raise TextExtractionError("PDF is encrypted; cannot extract text.")
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(parts)

    # ----- DOCX -----

    def _extract_docx(self, file_bytes: bytes) -> str:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            raise TextExtractionError(f"Could not open DOCX: {e}") from e

        parts: list[str] = []

        # Paragraphs in the document body
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables — many resumes use tables for layout
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)

    # ----- Cleanup -----

    def _clean(self, text: str) -> str:
        """Remove obvious noise so we send fewer tokens to the LLM."""
        if not text:
            return ""

        lines = text.splitlines()
        cleaned_lines: list[str] = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            # Drop lines that are just a page number
            if _PAGE_NUMBER_PATTERN.match(stripped):
                continue
            # Collapse internal whitespace runs
            stripped = re.sub(r"[ \t]+", " ", stripped)
            cleaned_lines.append(stripped)

        # Collapse multiple consecutive blank lines (already removed above,
        # but in case the join introduces any)
        result = "\n".join(cleaned_lines)
        result = re.sub(r"\n{3,}", "\n\n", result)
        return result.strip()
